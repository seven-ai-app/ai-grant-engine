"""Generate Hebrew RTL Word document for Tnufa grant application.

Produces a fully structured document including:
- Cover info table
- All 21 application sections as proper headings
- R&D tasks table
- Competitor table
- Capability gap table
- Budget table
Target: professional, ~25 pages max.
"""

import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..graph.state import GrantState
from ..tools.hebrew_utils import format_currency

logger = logging.getLogger(__name__)

# Section IDs that have dedicated tables (content already embedded in table)
STRUCTURED_TABLE_SECTIONS = {"rd_tasks", "competition"}


def generate_application_docx(state: GrantState, output_dir: Path) -> Path:
    """Generate the full Tnufa grant application as a Hebrew RTL Word document."""
    doc = Document()
    _set_document_margins(doc)

    # ------------------------------------------------------------------ #
    # Cover page
    # ------------------------------------------------------------------ #
    _add_cover_page(doc, state)

    # ------------------------------------------------------------------ #
    # Table of contents (manual)
    # ------------------------------------------------------------------ #
    doc.add_page_break()
    _add_heading(doc, "תוכן עניינים", level=1)
    for section in state.get("sections", []):
        _add_paragraph(doc, f"• {section['title_he']}", size=11)
    _add_paragraph(doc, "• טבלת משימות מו\"פ", size=11)
    _add_paragraph(doc, "• טבלת יכולות - פערים ויעדים", size=11)
    _add_paragraph(doc, "• טבלת מתחרים", size=11)
    _add_paragraph(doc, "• סיכום תקציבי", size=11)

    # ------------------------------------------------------------------ #
    # Application sections
    # ------------------------------------------------------------------ #
    doc.add_page_break()
    for section in state.get("sections", []):
        _add_heading(doc, section["title_he"], level=1)

        content = section.get("content_he", "")
        if content:
            # Split on double newlines for multi-paragraph content
            paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
            if not paragraphs:
                paragraphs = [content]
            for para_text in paragraphs:
                _add_paragraph(doc, para_text, size=11)
        doc.add_paragraph()  # spacing

    # ------------------------------------------------------------------ #
    # R&D tasks table
    # ------------------------------------------------------------------ #
    rd_tasks = state.get("rd_tasks", [])
    if rd_tasks:
        doc.add_page_break()
        _add_heading(doc, "טבלת משימות המו\"פ", level=1)
        _add_rd_tasks_table(doc, rd_tasks)

    # ------------------------------------------------------------------ #
    # Capability gap table
    # ------------------------------------------------------------------ #
    cap_table = state.get("capability_table", [])
    if cap_table:
        doc.add_paragraph()
        _add_heading(doc, "טבלת יכולות המיזם - פערים ויעדים", level=1)
        _add_capability_table(doc, cap_table)

    # ------------------------------------------------------------------ #
    # Competitor table
    # ------------------------------------------------------------------ #
    competitor_table = state.get("competitor_table", [])
    if competitor_table:
        doc.add_paragraph()
        _add_heading(doc, "טבלת מתחרים", level=1)
        _add_competitor_table(doc, competitor_table)

    # ------------------------------------------------------------------ #
    # Budget summary
    # ------------------------------------------------------------------ #
    if state.get("budget_lines"):
        doc.add_page_break()
        _add_heading(doc, "סיכום תקציבי", level=1)
        _add_budget_table(doc, state)

    # ------------------------------------------------------------------ #
    # Save
    # ------------------------------------------------------------------ #
    filename = f"בקשת_מענק_תנופה_{state['startup_name'].replace(' ', '_')}.docx"
    output_path = output_dir / filename
    doc.save(str(output_path))
    logger.info(f"Generated DOCX: {output_path}")
    return output_path


# ====================================================================== #
# Cover page
# ====================================================================== #

def _add_cover_page(doc: Document, state: GrantState):
    """Add a structured cover page with info table."""
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = _add_paragraph(doc, "בקשה לתמיכה במסלול תנופה (Ideation)", size=20, bold=True)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = _add_paragraph(doc, "רשות החדשנות הישראלית", size=14)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Info table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    info_rows = [
        ("שם המיזם", state["startup_name"]),
        ("סוג מגיש", "יזם פרטי" if state.get("entity_type") == "private_entrepreneur" else "חברה בע\"מ"),
        ("תיאור קצר", state.get("startup_description", "")[:200]),
        ("תקציב מבוקש (ש\"ח)", format_currency(state.get("total_budget", 0))),
        ("מענק מבוקש (ש\"ח)", format_currency(state.get("grant_amount", 0))),
        ("השתתפות עצמית (ש\"ח)", format_currency(state.get("self_funding", 0))),
    ]

    # Remove the empty first row and build properly
    for label, value in info_rows:
        row = table.add_row()
        _set_cell_text(row.cells[0], label, bold=True)
        _set_cell_text(row.cells[1], str(value))

    # Remove the initial empty row
    tbl = table._tbl
    first_tr = tbl.tr_lst[0]
    tbl.remove(first_tr)


# ====================================================================== #
# R&D Tasks table
# ====================================================================== #

def _add_rd_tasks_table(doc: Document, rd_tasks: list[dict]):
    """Add a structured R&D tasks table."""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = ["#", "שם המשימה", "פירוט שלבים", "התחלה", "סיום"]
    for i, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], header, bold=True)

    for idx, task in enumerate(rd_tasks, 1):
        row = table.add_row()
        _set_cell_text(row.cells[0], str(idx))
        _set_cell_text(row.cells[1], task.get("task_name", ""))
        _set_cell_text(row.cells[2], task.get("description", ""))
        _set_cell_text(row.cells[3], task.get("start_mm_yy", ""))
        _set_cell_text(row.cells[4], task.get("end_mm_yy", ""))


# ====================================================================== #
# Capability table
# ====================================================================== #

def _add_capability_table(doc: Document, cap_table: list[dict]):
    """Add a capability gap table."""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    headers = ["יכולת", "מצב נוכחי", "מצב בסיום הפרויקט"]
    for i, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], header, bold=True)

    for item in cap_table:
        row = table.add_row()
        _set_cell_text(row.cells[0], item.get("capability", ""))
        _set_cell_text(row.cells[1], item.get("current_state", ""))
        _set_cell_text(row.cells[2], item.get("target_state", ""))


# ====================================================================== #
# Competitor table
# ====================================================================== #

def _add_competitor_table(doc: Document, competitors: list[dict]):
    """Add a competitor comparison table."""
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = ["#", "שם החברה", "אתר", "מאפיינים", "מחיר", "נתח שוק"]
    for i, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], header, bold=True)

    for idx, comp in enumerate(competitors, 1):
        row = table.add_row()
        _set_cell_text(row.cells[0], str(idx))
        _set_cell_text(row.cells[1], comp.get("name", ""))
        _set_cell_text(row.cells[2], comp.get("url", ""))
        _set_cell_text(row.cells[3], comp.get("features", ""))
        _set_cell_text(row.cells[4], comp.get("price", ""))
        _set_cell_text(row.cells[5], comp.get("market_share", ""))


# ====================================================================== #
# Budget table
# ====================================================================== #

def _add_budget_table(doc: Document, state: GrantState):
    """Add a budget summary table with totals."""
    budget_lines = state.get("budget_lines", [])
    if not budget_lines:
        return

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    headers = ["תיאור", "קטגוריה", "סכום כולל (ש\"ח)", "מענק (ש\"ח)", "נימוק"]
    for i, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], header, bold=True)

    # Category display names
    cat_names = {
        "subcontractors": "קבלני משנה",
        "materials": "חומרים ורכיבים",
        "ip_patents": "קניין רוחני",
        "business_development": "פיתוח עסקי",
        "equipment_depreciation": "פחת ציוד",
        "travel_abroad": "נסיעות לחו\"ל",
    }

    for line in budget_lines:
        row = table.add_row()
        _set_cell_text(row.cells[0], line.get("description_he", ""))
        cat_key = line.get("category", "")
        _set_cell_text(row.cells[1], cat_names.get(cat_key, cat_key))
        _set_cell_text(row.cells[2], f"{line.get('amount', 0):,.0f}")
        _set_cell_text(row.cells[3], f"{line.get('grant_portion', 0):,.0f}")
        justification = line.get("justification_he", "")
        _set_cell_text(row.cells[4], justification[:120])

    # Totals row
    total_row = table.add_row()
    _set_cell_text(total_row.cells[0], "סה\"כ", bold=True)
    _set_cell_text(total_row.cells[1], "", bold=True)
    _set_cell_text(total_row.cells[2], f"{state.get('total_budget', 0):,.0f}", bold=True)
    _set_cell_text(total_row.cells[3], f"{state.get('grant_amount', 0):,.0f}", bold=True)
    _set_cell_text(total_row.cells[4], "")


# ====================================================================== #
# Low-level formatting helpers
# ====================================================================== #

def _set_document_margins(doc: Document):
    """Set page margins for a professional look."""
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add an RTL heading."""
    heading = doc.add_heading(text, level=level)
    _make_rtl(heading)
    heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return heading


def _add_paragraph(doc: Document, text: str, size: int = 11, bold: bool = False):
    """Add an RTL paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "David"
    _make_rtl(para)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return para


def _set_cell_text(cell, text: str, bold: bool = False):
    """Set cell text with RTL alignment."""
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "David"
    run.font.size = Pt(10)
    _make_rtl(para)
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _make_rtl(element):
    """Add RTL bidi mark to a paragraph XML element."""
    try:
        pPr = element._element.get_or_add_pPr()
        # Avoid duplicate bidi elements
        existing = pPr.findall(qn("w:bidi"))
        if not existing:
            bidi = OxmlElement("w:bidi")
            pPr.append(bidi)
    except Exception:
        pass
